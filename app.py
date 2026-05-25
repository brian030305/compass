# python -m streamlit run app.py
import streamlit as st
import pandas as pd
import google.generativeai as genai
import io
import re
from docx import Document
from datetime import datetime, timedelta
from streamlit_gsheets import GSheetsConnection
from data_manager import (
    fetch_safety_cert_data, fetch_mss_data, fetch_ktl_data,
    fetch_kiat_data, fetch_keit_min_data, fetch_keit_rd_data,
    fetch_national_business_api, fetch_local_keit_announcement, 
    fetch_mss_tech_cert_api, fetch_bizinfo_api
)

# 1. 기본 설정
st.set_page_config(page_title="창업나침반 - Startup Compass", layout="wide")

st.markdown(
    """
    <div style="background-color:#0e243a; padding:15px; border-radius:10px; margin-bottom:25px;">
        <h2 style="color:white; margin:0; display:inline-block;">🧭 창업나침반 <span style="font-size:16px; font-weight:normal; color:#b0c4de;">Startup Compass</span></h2>
    </div>
    """, 
    unsafe_allow_html=True
)

if 'current_page' not in st.session_state:
    st.session_state.current_page = st.query_params.get("page", "대시보드")

def change_page(page_name):
    st.session_state.current_page = page_name
    st.query_params["page"] = page_name

# ==========================================
# 🚨 로그인 시스템 및 차단기 (구글 시트 실시간 연동)
# ==========================================
# 🎯 구글 시트 연결 객체 생성
conn = st.connection("gsheets", type=GSheetsConnection)

if st.query_params.get("auth") == "success":
    st.session_state["logged_in"] = True
    
    # 새로고침 시 URL 파라미터를 통해 세션 복구
    if 'company_name' not in st.session_state:
        st.session_state.user_id = st.query_params.get("uid", "admin")
        st.session_state.company_name = st.query_params.get("company", "")
        st.session_state.industry = st.query_params.get("industry_val", "")
        st.session_state.tech_field = st.query_params.get("tech", "")
        st.session_state.location = st.query_params.get("location", "")
        st.session_state.dashboard_metrics = None

if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if not st.session_state["logged_in"]:
    login_container = st.empty()
    with login_container.container():
        st.title("🔒 창업나침반 - 로그인")
        st.info("서비스를 이용하시려면 로그인이 필요합니다.")
        
        with st.form(key="login_form"):
            user_id = st.text_input("아이디")
            user_pw = st.text_input("비밀번호", type="password")
            submit_button = st.form_submit_button(label="로그인", type="primary", use_container_width=True)
            
        if submit_button:
            with st.spinner("구글 클라우드 서버에서 유저 정보를 확인 중입니다..."):
                try:
                    # 🎯 구글 시트에서 데이터 읽어오기
                    users_df = conn.read()
                    
                    # 아이디가 일치하는 행(Row) 찾기
                    user_row = users_df[users_df['id'].astype(str) == user_id]
                    
                    # 계정이 존재하고 비밀번호가 일치하는지 검증
                    if not user_row.empty and str(user_row.iloc[0]['pw']) == user_pw:
                        st.session_state["logged_in"] = True
                        st.query_params["auth"] = "success"
                        st.query_params["uid"] = user_id
                        
                        # 🎯 구글 시트에 적혀있던 데이터를 세션에 장착
                        st.session_state.user_id = user_id
                        st.session_state.company_name = str(user_row.iloc[0]['company'])
                        st.session_state.location = str(user_row.iloc[0]['location'])
                        st.session_state.industry = str(user_row.iloc[0]['industry'])
                        st.session_state.tech_field = str(user_row.iloc[0]['tech'])
                        
                        st.query_params["company"] = st.session_state.company_name
                        st.query_params["industry_val"] = st.session_state.industry
                        st.query_params["tech"] = st.session_state.tech_field
                        st.query_params["location"] = st.session_state.location
                        
                        st.session_state.dashboard_metrics = None
                        login_container.empty() 
                        st.rerun() 
                    else:
                        st.error("아이디 또는 비밀번호가 일치하지 않습니다.")
                except Exception as e:
                    st.error("데이터베이스(구글 시트) 접근 중 오류가 발생했습니다. 권한 설정이나 키 값을 확인해주세요.")
                
    st.stop()

# --- AI 스마트 분류 함수 ---
@st.cache_data(ttl=3600) 
def get_ai_classified_data():
    df = get_ai_classified_data()
    if df.empty: return df
    
    classification_prompt = """
    다음 지원사업 공고 제목을 읽고, 아래 5개 업종 중 가장 적합한 하나를 선택해. 
    답변은 딱 업종명만 출력할 것.
    업종 리스트: ['IT/소프트웨어', '제조업', '바이오/헬스케어', '에너지/환경', '기타']
    공고명: {title}
    """
    
    categories = []
    progress_bar = st.progress(0)
    for i, (_, row) in enumerate(df.iterrows()):
        title = str(row.get('사업명', '')) 
        response = st.session_state.chat_session.send_message(classification_prompt.format(title=title))
        categories.append(response.text.strip())
        progress_bar.progress((i + 1) / len(df))
    
    df['업종태그'] = categories
    progress_bar.empty()
    return df


# ==========================================
# 🎯 기업 정보 수정 팝업창 함수
# ==========================================
@st.dialog("⚙️ 기업 정보 수정")
def edit_company_profile():
    st.write("변경하실 정보를 입력해 주세요. 저장 시 대시보드 데이터가 실시간으로 변경됩니다.")
    
    current_company = st.session_state.get('company_name', '테크스타트업(주)')
    current_industry = st.session_state.get('industry', '선택해주세요')
    current_tech = st.session_state.get('tech_field', '')
    current_location = st.session_state.get('location', '전국')
    
    # 🎯 [핵심] 입력칸들을 form으로 묶어 엔터키 제출을 활성화합니다.
    with st.form(key="edit_company_form"):
        company_input = st.text_input("기업명", value=current_company)
        
        location_options = ["전국", "서울", "경기", "인천", "부산", "대구", "대전", "광주", "울산", "세종", "강원", "충북", "충남", "전북", "전남", "경북", "경남", "제주"]
        loc_index = location_options.index(current_location) if current_location in location_options else 0
        location_input = st.selectbox("기업 소재지(지역)", location_options, index=loc_index)
        
        industry_options = ["선택해주세요", "IT/소프트웨어", "제조업", "바이오/헬스케어", "에너지/환경", "기타"]
        ind_index = industry_options.index(current_industry) if current_industry in industry_options else 0
        industry_input = st.selectbox("어떤 업종에 속하시나요?", industry_options, index=ind_index)
        
        tech_field_input = st.text_input("필요한 기술 분야 키워드", value=current_tech, placeholder="예: 드론, 인공지능")
        
        st.markdown("---")
        # st.button 대신 st.form_submit_button을 사용해야 엔터키와 연동됩니다.
        submit_btn = st.form_submit_button("💾 변경사항 저장 및 실시간 반영", type="primary", use_container_width=True)
        
    if submit_btn:
        st.session_state.company_name = company_input
        st.session_state.industry = industry_input
        st.session_state.tech_field = tech_field_input
        st.session_state.location = location_input
        
        st.query_params["company"] = company_input
        st.query_params["industry_val"] = industry_input
        st.query_params["tech"] = tech_field_input
        st.query_params["location"] = location_input
        
        # 🎯 [핵심] 수정한 정보를 구글 시트에 실시간으로 덮어쓰기
        with st.spinner("변경 사항을 구글 클라우드 DB에 영구 저장 중입니다..."):
            try:
                users_df = conn.read()
                # 현재 로그인된 아이디의 줄(Row)을 찾아내어 내용물 교체
                mask = users_df['id'].astype(str) == st.session_state.user_id
                if mask.any():
                    users_df.loc[mask, 'company'] = company_input
                    users_df.loc[mask, 'location'] = location_input
                    users_df.loc[mask, 'industry'] = industry_input
                    users_df.loc[mask, 'tech'] = tech_field_input
                    
                    # 덮어쓴 데이터를 다시 구글 시트에 전송하여 업데이트
                    conn.update(data=users_df)
            except Exception as e:
                st.error(f"구글 시트 업데이트 중 오류 발생: {e}")
        
        st.session_state.dashboard_metrics = None 
        st.session_state.survival_report = None
        if "chat_session" in st.session_state:
            del st.session_state.chat_session
            
        st.success("기업 정보가 성공적으로 변경 및 저장되었습니다!")
        st.rerun()

# 2. 사이드바 메뉴
with st.sidebar:
    st.markdown(f"**환영합니다, {st.session_state.get('company_name', '고객')}님!**")
    if st.button("🚪 로그아웃", use_container_width=True):
        st.session_state["logged_in"] = False
        # 인증 정보 삭제
        st.query_params.clear() 
        st.rerun()
    if st.button("⚙️ 기업 정보 수정", use_container_width=True, type="secondary"):
        edit_company_profile()

    st.divider()
    st.markdown("### 🛠️ 메뉴")
    st.button("📊 대시보드", use_container_width=True, type="primary" if st.session_state.current_page == '대시보드' else "secondary", on_click=change_page, args=('대시보드',))
    st.button("💬 AI 어시스턴트", use_container_width=True, type="primary" if st.session_state.current_page == 'AI 어시스턴트' else "secondary", on_click=change_page, args=('AI 어시스턴트',))
    st.button("✨ AI 매칭", use_container_width=True, type="primary" if st.session_state.current_page == 'AI 매칭' else "secondary", on_click=change_page, args=('AI 매칭',))
    st.button("📈 생존율 예측", use_container_width=True, type="primary" if st.session_state.current_page == '생존율 예측' else "secondary", on_click=change_page, args=('생존율 예측',))
    st.button("📅 지원 캘린더", use_container_width=True, type="primary" if st.session_state.current_page == '지원 캘린더' else "secondary", on_click=change_page, args=('지원 캘린더',))
    st.button("📄 보고서 생성", use_container_width=True, type="primary" if st.session_state.current_page == '보고서 생성' else "secondary", on_click=change_page, args=('보고서 생성',))
    
    st.divider()
   

# 3. AI 및 데이터 처리 로직 (캐싱 적용)
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

def search_safety_cert(): return fetch_safety_cert_data().to_string() if not fetch_safety_cert_data().empty else "데이터 없음"
def search_mss_support(): return fetch_mss_data().to_string() if not fetch_mss_data().empty else "데이터 없음"
def search_ktl_test(): return fetch_ktl_data().to_string() if not fetch_ktl_data().empty else "데이터 없음"
def search_kiat_worldclass(): return fetch_kiat_data().to_string() if not fetch_kiat_data().empty else "데이터 없음"
def search_keit_ministry(): return fetch_keit_min_data().to_string() if not fetch_keit_min_data().empty else "데이터 없음"
def search_keit_rnd(): return fetch_keit_rd_data().to_string() if not fetch_keit_rd_data().empty else "데이터 없음"

tools_list = [search_safety_cert, search_mss_support, search_ktl_test, search_kiat_worldclass, search_keit_ministry, search_keit_rnd]

@st.cache_data(ttl=600)
def get_ai_classified_data():
    df_mss = fetch_mss_data()
    df_rd = fetch_keit_rd_data()
    valid_dfs = [d for d in [df_mss, df_rd] if not d.empty]
    return pd.concat(valid_dfs, ignore_index=True) if valid_dfs else pd.DataFrame()

base_instruction = "당신은 한국의 스타트업과 중소기업을 돕는 최고 수준의 AI 컨설턴트입니다."
if "industry" in st.session_state and "tech_field" in st.session_state:
    if st.session_state.industry != "선택해주세요" or st.session_state.tech_field:
        base_instruction += f"\n\n[고객 정보]\n- 업종: {st.session_state.industry}\n- 기술 분야: {st.session_state.tech_field}\n★이 기업 정보를 기준으로 사업을 필터링하세요."

if "chat_session" not in st.session_state:
    model = genai.GenerativeModel(model_name="gemini-2.5-flash", tools=tools_list, system_instruction=base_instruction)
    st.session_state.chat_session = model.start_chat(enable_automatic_function_calling=True)

# 4. 메인 화면 출력부
company_name = st.session_state.get('company_name', st.query_params.get("company", "테크스타트업(주)"))
ind_str = st.session_state.get('industry', st.query_params.get("industry", "선택해주세요"))
tech_str = st.session_state.get('tech_field', st.query_params.get("tech", ""))
loc_str = st.session_state.get('location', "전국")

# ------- 대시보드 설정 -----------
if st.session_state.current_page == '대시보드':
    if 'show_chatbot' not in st.session_state:
        st.session_state.show_chatbot = False

    header_col1, header_col2 = st.columns([8, 2])
    with header_col1:
        st.header(f"안녕하세요, {company_name}")
        st.caption(f"💡 {ind_str} · 핵심 기술: {tech_str if tech_str else '미설정'}")
    with header_col2:
        if st.button("💬 AI 챗봇 켜기/끄기", use_container_width=True):
            st.session_state.show_chatbot = not st.session_state.show_chatbot
            st.rerun()
            
    if st.session_state.show_chatbot:
        main_col, ai_col = st.columns([7, 3], gap="large") 
    else:
        main_col, ai_col = st.columns([1, 0.01]) 

    with main_col:
        is_info_set = (ind_str != "선택해주세요" and ind_str != "")
        
        if is_info_set:
            # 🎯 [핵심 1] AI 매칭 탭과 동일한 실시간 API 호출 및 완벽 필터링 적용
            df = fetch_bizinfo_api()
            
            if not df.empty and 'pblancNm' in df.columns:
                df = df.rename(columns={'pblancNm': '사업명', 'pancInsttNm': '소관기관', 'reqstBeginDe': '접수시작일', 'reqstEndDe': '마감일', 'areaNm': '지역'})
                
                if '마감일' not in df.columns: df['마감일'] = ''
                if '접수시작일' not in df.columns: df['접수시작일'] = ''
                
                def rescue_dates(r):
                    e_val = str(r.get('마감일', ''))
                    if pd.isna(pd.to_datetime(e_val, errors='coerce')):
                        row_str = " ".join([str(v) for v in r.values if pd.notna(v)])
                        dates = re.findall(r'20\d{2}[-./]\d{2}[-./]\d{2}', row_str)
                        if dates:
                            dates = sorted([d.replace('.', '-').replace('/', '-') for d in dates])
                            return dates[0], dates[-1]
                        return '', ''
                    return str(r.get('접수시작일', '')), e_val
                df[['접수시작일', '마감일']] = df.apply(lambda x: pd.Series(rescue_dates(x)), axis=1)

                if loc_str != "전국":
                    if '지역' not in df.columns: df['지역'] = ''
                    if '소관기관' not in df.columns: df['소관기관'] = ''
                    if '사업명' not in df.columns: df['사업명'] = ''
                    df['검색용_텍스트'] = df['지역'].astype(str).fillna('') + ' ' + df['소관기관'].astype(str).fillna('') + ' ' + df['사업명'].astype(str).fillna('')
                    all_regions = ["서울", "경기", "인천", "부산", "대구", "대전", "광주", "울산", "세종", "강원", "충북", "충남", "전북", "전남", "경북", "경남", "제주"]
                    other_regions = [r for r in all_regions if r != loc_str]
                    mask = ~df['검색용_텍스트'].str.contains('|'.join(other_regions), case=False, na=False) | df['검색용_텍스트'].str.contains(loc_str, case=False, na=False)
                    df = df[mask].drop(columns=['검색용_텍스트'])
                
                industry_keywords = {
                    "IT/소프트웨어": ["IT", "소프트웨어", "SW", "정보통신", "플랫폼", "앱", "데이터", "인공지능", "AI"],
                    "제조업": ["제조", "생산", "설비", "부품", "소재", "하드웨어", "공장", "가공"],
                    "바이오/헬스케어": ["바이오", "의료", "헬스", "건강", "의약", "생명", "제약", "병원"],
                    "에너지/환경": ["에너지", "환경", "녹색", "친환경", "탄소", "수소", "재생", "에코"],
                    "기타": []
                }
                if ind_str in industry_keywords and industry_keywords[ind_str]:
                    pattern = '|'.join(industry_keywords[ind_str])
                    df = df[df.apply(lambda row: row.astype(str).str.contains(pattern, case=False).any(), axis=1)]
                if tech_str:
                    df = df[df.apply(lambda row: row.astype(str).str.contains(tech_str, case=False).any(), axis=1)]
            
            # 🎯 [핵심 2] 진짜 리얼 데이터 기반 KPI 계산 로직
            match_count = len(df) if not df.empty else 0
            
            urgent_count = 0
            this_month_count = 0
            total_fund_calc = 0
            
            if match_count > 0:
                today = datetime.now().date()
                current_month = today.month
                current_year = today.year
                
                for _, row in df.iterrows():
                    # 1) 실제 날짜 기반 마감 임박 스캔
                    end_date_str = str(row.get('마감일', ''))
                    parsed_end = pd.to_datetime(end_date_str, errors='coerce')
                    
                    if pd.notna(parsed_end):
                        end_date_obj = parsed_end.date()
                        days_left = (end_date_obj - today).days
                        if 0 <= days_left <= 7:
                            urgent_count += 1
                        if end_date_obj.month == current_month and end_date_obj.year == current_year:
                            this_month_count += 1
                            
                    # 2) 텍스트 기반 실제 지원 금액 추출 (억, 천만 단위 정규식 스캔)
                    title_text = str(row.get('사업명', ''))
                    ok_match = re.search(r'(\d+(?:\.\d+)?)억', title_text)
                    if ok_match: total_fund_calc += float(ok_match.group(1))
                    
                    cheon_match = re.search(r'(\d+(?:\.\d+)?)천만', title_text)
                    if cheon_match: total_fund_calc += float(cheon_match.group(1)) * 0.1

            # --- [기능 1] 마감 임박 공지 (데일리 리포트 배너) ---
            if urgent_count > 0:
                st.error(f"🚨 **주목!** 귀사에 적합한 공고 중 마감이 **7일 이내로 임박한 공고가 {urgent_count}건** 있습니다. 캘린더 탭에서 서둘러 확인해 보세요!")
            elif match_count > 0:
                st.info(f"💡 오늘 기준 **{ind_str}** 분야에 지원 가능한 새로운 맞춤형 공고가 **{match_count}건** 대기 중입니다.")
            else:
                st.info("💡 조건에 맞는 지원사업이 없습니다.")

            st.divider()
            
            # --- [기능 2] 대시보드 핵심 지표 요약표 (리얼 데이터) ---
            kpi1, kpi2, kpi3, kpi4 = st.columns(4)
            
            # 생존율은 AI 계산 결과가 있으면 평균값, 없으면 '분석 대기' 처리
            if st.session_state.get('dashboard_metrics'):
                real_survival_rate = sum(st.session_state.dashboard_metrics) / 4
                survival_display = f"{real_survival_rate:.1f}%"
            else:
                survival_display = "분석 전"
                
            fund_display = f"{total_fund_calc:.1f}억원" if total_fund_calc > 0 else "별도 공고 참조"
            
            with kpi1:
                st.metric(label="매칭된 지원사업", value=f"{match_count}건", delta="실시간 API 연동")
            with kpi2:
                st.metric(label="AI 예상 생존율", value=survival_display, delta="버튼을 눌러 분석" if survival_display == "분석 전" else "정밀 분석 완료")
            with kpi3:
                st.metric(label="이번 달 마감", value=f"{this_month_count}건", delta="신청 임박", delta_color="inverse")
            with kpi4:
                st.metric(label="총 지원 가능 금액", value=fund_display, delta="공고 텍스트 스캔" if total_fund_calc > 0 else "상세 내용 확인 필요")
                
            st.divider()
            
            # --- [기능 3] 생존율 리스크 시각화 바 ---
            if 'last_ind_str' not in st.session_state:
                st.session_state.last_ind_str = ind_str
            
            if st.session_state.last_ind_str != ind_str:
                st.session_state.dashboard_metrics = None
                st.session_state.last_ind_str = ind_str
                st.rerun()

            if st.session_state.get('dashboard_metrics') is None:
                st.info("🔄 최신 공고 데이터를 기반으로 기업 생존율과 세부 역량을 진단하려면 아래 버튼을 눌러주세요.")
                
                if st.button("🤖 AI 생존율 정밀 계산 시작", use_container_width=True):
                    with st.spinner("실시간 공고 텍스트를 분석하여 역량 점수를 도출 중입니다... (약 5~10초 소요)"):
                        data_summary = df.head(10)['사업명'].to_string() if not df.empty else "데이터 없음"
                        
                        # 🚨 [수정 적용 완료] 대화 내역 보호를 위해 백그라운드 전용 모델 사용
                        background_model = genai.GenerativeModel(model_name="gemini-2.5-flash")
                        analysis_prompt = f"""
                        업종: {ind_str}, 기술: {tech_str}
                        매칭된 실제 공고 목록: {data_summary}
                        이 공고들의 성격과 우리 기업의 키워드를 바탕으로 재무안정성, 기술경쟁력, 시장성장성, 팀역량을 0~100점으로 평가해줘.
                        반드시 숫자 4개만 콤마로 구분해서 줄 것. 예: 75,82,60,88
                        """
                        try:
                            res = background_model.generate_content(analysis_prompt)
                            found_numbers = re.findall(r'\d+', res.text)
                            
                            if len(found_numbers) >= 4:
                                st.session_state.dashboard_metrics = [int(n) for n in found_numbers[:4]]
                                st.rerun() 
                            else:
                                st.error("분석 실패: AI가 유효한 점수를 반환하지 못했습니다.")
                        except Exception as e:
                            st.error(f"AI 호출 중 오류가 발생했습니다: {e}")

            if st.session_state.dashboard_metrics:
                s1, s2, s3, s4 = st.session_state.dashboard_metrics
                risk_col1, risk_col2 = st.columns(2)
                with risk_col1:
                    st.markdown("**재무 안정성**")
                    st.progress(s1, text=f"{s1}%")
                    st.markdown("**시장 성장성**")
                    st.progress(s3, text=f"{s3}%")
                    
                with risk_col2:
                    st.markdown("**기술 경쟁력**")
                    st.progress(s2, text=f"{s2}%")
                    st.markdown("**팀 역량**")
                    st.progress(s4, text=f"{s4}%")
                    
        else:
            st.info("💡 좌측 사이드바에서 기업 정보를 설정하시면 맞춤형 지원사업 알림과 생존율 진단을 받아보실 수 있습니다.")

            
    # --- [기능 4] 우측 AI 챗봇 토글 영역 ---
    if st.session_state.show_chatbot:
        with ai_col:
            st.markdown("### 🟢 AI 어시스턴트")
            
            chat_container = st.container(height=550)
            with chat_container:
                for message in st.session_state.chat_session.history:
                    role = "user" if message.role == "user" else "assistant"
                    with st.chat_message(role):
                        st.markdown(message.parts[0].text)
            
            user_input = st.chat_input("질문하거나 지시를 내려보세요...", key="dashboard_chat_input")
            
            if user_input:
                with chat_container:
                    with st.chat_message("user"):
                        st.markdown(user_input)
                    
                    with st.chat_message("assistant"):
                        with st.spinner("AI가 분석 중..."):
                            response = st.session_state.chat_session.send_message(user_input)
                            st.markdown(response.text)
                st.rerun()

# ------- AI 어시스턴트 설정 -----------
elif st.session_state.current_page == 'AI 어시스턴트':
    st.header("💬 AI 어시스턴트")
    for message in st.session_state.chat_session.history:
        if message.role == "user" or (message.role == "model" and message.parts[0].text):
            with st.chat_message(message.role):
                st.markdown(message.parts[0].text)
    user_input = st.chat_input("궁금한 공공데이터에 대해 물어보세요!")
    if user_input:
        with st.chat_message("user"): st.markdown(user_input)
        with st.chat_message("model"):
            with st.spinner("AI가 데이터를 분석 중입니다..."):             
                response = st.session_state.chat_session.send_message(user_input, stream=True)                         
                st.write_stream(response)
        st.rerun()

# ------- AI 매칭 설정 -----------
elif st.session_state.current_page == 'AI 매칭':
    st.subheader("✨ AI 맞춤형 지원사업 매칭")
    st.caption(f"현재 매칭 조건 ➔ 지역: **{loc_str}** | 업종: **{ind_str}** | 기술 키워드: **{tech_str if tech_str else '미설정'}**")
    
    # 1. 🚀 리스트 불러오기 및 AI 추천 엔진 가동
    if st.button("🚀 실시간 공고 리스트 및 AI 추천 분석 불러오기", type="primary"):
        with st.spinner("서버에서 공고를 수집하고, AI가 귀사에 가장 적합한 사업을 정밀 분석 중입니다... (약 5~10초 소요)"):
            
            df = fetch_bizinfo_api()
            
            if df.empty:
                st.warning("서버에서 불러올 수 있는 실시간 공고가 없거나 API 키를 확인해주세요.")
                if 'matching_list_df' in st.session_state:
                    del st.session_state.matching_list_df
            else:
                if 'pblancNm' in df.columns:
                    df = df.rename(columns={
                        'pblancNm': '사업명',
                        'pancInsttNm': '소관기관',
                        'reqstBeginDe': '접수시작일',
                        'reqstEndDe': '마감일',
                        'pblancUrl': '상세링크',
                        'areaNm': '지역' 
                    })
                    
                    if '마감일' not in df.columns: df['마감일'] = ''
                    if '접수시작일' not in df.columns: df['접수시작일'] = ''
                    
                    def rescue_dates(r):
                        e_val = str(r.get('마감일', ''))
                        s_val = str(r.get('접수시작일', ''))
                        if pd.isna(pd.to_datetime(e_val, errors='coerce')):
                            row_str = " ".join([str(v) for v in r.values if pd.notna(v)])
                            dates = re.findall(r'20\d{2}[-./]\d{2}[-./]\d{2}', row_str)
                            if dates:
                                dates = sorted([d.replace('.', '-').replace('/', '-') for d in dates])
                                return dates[0], dates[-1]
                            return '', ''
                        return s_val, e_val
                        
                    df[['접수시작일', '마감일']] = df.apply(lambda x: pd.Series(rescue_dates(x)), axis=1)

                # 지역 필터링
                if loc_str != "전국":
                    if '지역' not in df.columns: df['지역'] = ''
                    if '소관기관' not in df.columns: df['소관기관'] = ''
                    if '사업명' not in df.columns: df['사업명'] = ''
                        
                    df['검색용_텍스트'] = df['지역'].astype(str).fillna('') + ' ' + df['소관기관'].astype(str).fillna('') + ' ' + df['사업명'].astype(str).fillna('')
                    all_regions = ["서울", "경기", "인천", "부산", "대구", "대전", "광주", "울산", "세종", "강원", "충북", "충남", "전북", "전남", "경북", "경남", "제주"]
                    other_regions = [r for r in all_regions if r != loc_str]
                    
                    mask = ~df['검색용_텍스트'].str.contains('|'.join(other_regions), case=False, na=False) | df['검색용_텍스트'].str.contains(loc_str, case=False, na=False)
                    df = df[mask].drop(columns=['검색용_텍스트'])
                
                # 업종/기술 필터링
                industry_keywords = {
                    "IT/소프트웨어": ["IT", "소프트웨어", "SW", "정보통신", "플랫폼", "앱", "데이터", "인공지능", "AI"],
                    "제조업": ["제조", "생산", "설비", "부품", "소재", "하드웨어", "공장", "가공"],
                    "바이오/헬스케어": ["바이오", "의료", "헬스", "건강", "의약", "생명", "제약", "병원"],
                    "에너지/환경": ["에너지", "환경", "녹색", "친환경", "탄소", "수소", "재생", "에코"],
                    "기타": []
                }
                if ind_str != "선택해주세요" and ind_str in industry_keywords:
                    keywords = industry_keywords[ind_str]
                    if keywords:
                        pattern = '|'.join(keywords)
                        df = df[df.apply(lambda row: row.astype(str).str.contains(pattern, case=False).any(), axis=1)]
                
                if tech_str:
                    df = df[df.apply(lambda row: row.astype(str).str.contains(tech_str, case=False).any(), axis=1)]
                
                if df.empty:
                    st.info(f"선택하신 지역({loc_str}), 업종 및 기술 키워드와 매칭되는 실시간 공고가 없습니다.")
                    if 'matching_list_df' in st.session_state:
                        del st.session_state.matching_list_df
                else:
                    df['관심 등록'] = False
                    df['구분'] = "🔵 일반 매칭" # 기본 상태 설정
                    
                    # 🎯 [핵심] 실제 Gemini AI를 호출하여 최대 5개의 찰떡 공고를 픽(Pick)합니다.
                    try:
                        project_titles = df['사업명'].tolist()
                        # 토큰 절약 및 속도 향상을 위해 상위 50개 리스트만 AI에게 전달
                        titles_text = "\n".join([f"- {t}" for t in project_titles[:50]]) 
                        
                        ai_prompt = f"""
                        우리 기업은 '{loc_str}' 지역의 '{ind_str}' 업종이며, 핵심 기술은 '{tech_str}'입니다.
                        아래 지원사업 목록 중에서 우리 기업에게 가장 강력하게 추천할 만한 핵심 사업을 최대 5개만 골라주세요.
                        답변은 반드시 사업명 텍스트만 한 줄에 하나씩 출력하고, 부연 설명은 절대 하지 마세요.
                        [목록]
                        {titles_text}
                        """
                        # 🚨 [핵심 수정] 챗봇 세션 대신, 백그라운드 전용 일회성 모델을 즉석 생성하여 질문합니다.
                        background_model = genai.GenerativeModel(model_name="gemini-2.5-flash")
                        ai_response = background_model.generate_content(ai_prompt)
                        
                        recommended_titles = [t.strip().replace('- ', '').strip() for t in ai_response.text.split('\n') if t.strip()]
                        
                        for rec_title in recommended_titles:
                            mask = df['사업명'].str.contains(re.escape(rec_title), case=False, na=False)
                            df.loc[mask, '구분'] = "⭐ AI 강력 추천"
                            
                    except Exception as e:
                        pass # AI 분석이 뻗더라도 일반 리스트는 정상 출력되도록 패스 처리
                        
                    # 🎯 파이썬 카테고리 기능을 이용해 '⭐ AI 강력 추천'이 무조건 1순위(최상단)가 되도록 강제 지정
                    df['구분'] = pd.Categorical(df['구분'], categories=["⭐ AI 강력 추천", "🔵 일반 매칭"], ordered=True)
                    df = df.sort_values(by='구분').reset_index(drop=True)
                    
                    st.session_state.matching_list_df = df.copy()
                    st.success(f"매칭된 공고 {len(df)}건 중, AI가 기업 정보 기반으로 가장 적합한 공고를 선별했습니다!")

    # 2. 🎯 화면 출력 구역 및 사용자 자율 검색창
    if 'matching_list_df' in st.session_state and not st.session_state.matching_list_df.empty:
        df_to_edit = st.session_state.matching_list_df.copy()
        
        # 🔍 [핵심 신규 기능] 공고 리스트 내에서 사용자 마음대로 검색 가능 (일부 단어만 쳐도 됨)
        search_keyword = st.text_input("🔍 공고 리스트 내 검색 (사업명 일부만 입력해도 실시간 필터링됩니다):", placeholder="예: 바우처, 청년, R&D, 지원 등")
        
        if search_keyword:
            # 검색어가 있으면 사업명 컬럼에서 해당 단어가 포함된 줄만 남김
            mask = df_to_edit['사업명'].str.contains(search_keyword, case=False, na=False)
            df_to_edit = df_to_edit[mask]
            
        if df_to_edit.empty:
            st.warning("입력하신 검색어와 일치하는 공고가 현재 매칭된 리스트에 없습니다.")
        else:
            with st.form(key="matching_selection_form"):
                edited_df = st.data_editor(
                    df_to_edit,
                    use_container_width=True,
                    column_order=['관심 등록', '구분', '사업명', '소관기관', '지역', '마감일', '상세링크'],
                    column_config={
                        "관심 등록": st.column_config.CheckboxColumn("선택", default=False),
                        "구분": st.column_config.TextColumn("분류", width="small"),
                        "사업명": st.column_config.TextColumn("사업명", width="large"),
                        "지역": st.column_config.TextColumn("지역", width="small"),
                        "마감일": st.column_config.TextColumn("마감일", width="small"),
                        "상세링크": st.column_config.LinkColumn("상세링크", display_text="🔗 공고 확인", width="small")
                    },
                    hide_index=True,
                    key="matching_table_editor"
                )
                
                st.markdown(" ")
                submit_calendar = st.form_submit_button("📅 선택한 공고를 지원 캘린더에 저장하기", type="primary", use_container_width=True)
            
            if submit_calendar:
                # 검색 결과 창에서 조작한 결과를 원본 메모리에 업데이트 (인덱스 추적)
                st.session_state.matching_list_df.update(edited_df)
                selected_rows = st.session_state.matching_list_df[st.session_state.matching_list_df['관심 등록'] == True]
                
                if selected_rows.empty:
                    st.warning("캘린더에 추가할 공고를 목록에서 먼저 선택(체크)해 주세요.")
                else:
                    if 'calendar_events' not in st.session_state:
                        st.session_state.calendar_events = []
                        
                    added_count = 0
                    for _, row in selected_rows.iterrows():
                        full_title = str(row.get('사업명', '이름 없음'))
                        if any(e['full_title'] == full_title for e in st.session_state.calendar_events):
                            continue
                            
                        raw_start = str(row.get('접수시작일', '')).strip()[:10]
                        raw_end = str(row.get('마감일', '')).strip()[:10]
                        
                        parsed_end = pd.to_datetime(raw_end, errors='coerce')
                        parsed_start = pd.to_datetime(raw_start, errors='coerce')
                        
                        if pd.notna(parsed_end):
                            end_date_str = parsed_end.strftime("%Y-%m-%d")
                        else:
                            end_date_str = datetime.now().strftime("%Y-%m-%d")
                            full_title = "[마감일 미정/상시] " + full_title
                            
                        if pd.notna(parsed_start):
                            start_date_str = parsed_start.strftime("%Y-%m-%d")
                        else:
                            start_date_str = end_date_str
                        
                        try:
                            end_date_obj = datetime.strptime(end_date_str, "%Y-%m-%d").date() + timedelta(days=1)
                            calc_end_str = end_date_obj.strftime("%Y-%m-%d")
                        except Exception:
                            calc_end_str = end_date_str
                        
                        st.session_state.calendar_events.append({
                            "title": full_title[:25] + "..." if len(full_title) > 25 else full_title,
                            "full_title": full_title,
                            "start": start_date_str,
                            "end": calc_end_str,
                            "url": str(row.get('상세링크', ''))
                        })
                        added_count += 1
                        
                    st.success(f"🎉 체크하신 {added_count}건의 공고가 나만의 지원 캘린더에 정상 등록되었습니다! '지원 캘린더' 메뉴에서 확인하세요.")


# ------- 생존율 예측 설정 -----------
elif st.session_state.current_page == '생존율 예측':
    st.subheader("📈 산업군 기반 생존율 정밀 진단 (빅데이터 연계)")
    if 'survival_report' not in st.session_state:
        st.session_state.survival_report = None
        
    if st.button("🚀 빅데이터 기반 생존율 정밀 분석 시작", type="primary"):
        if ind_str == "선택해주세요":
            st.warning("왼쪽 사이드바에서 기업 정보(업종 등)를 먼저 설정해주세요!")
        else:
            with st.spinner("과거 통계 데이터와 기업 데이터를 융합 분석 중입니다... (약 10~15초 소요)"):
                
                # 1. 과거 통계 및 실시간 API 융합
                df_biz = fetch_national_business_api()     
                df_keit = fetch_local_keit_announcement()  
                df_mss_sup = fetch_mss_data()              
                df_cert = fetch_mss_tech_cert_api()        
                
                # 2. 과도한 토큰 소모를 막기 위해, 현재 기업의 업종/기술 키워드와 일치하는 데이터만 일부 추출
                def get_summary(df, keyword, max_rows=3):
                    if df is None or df.empty: 
                        return "해당 데이터 없음"
                    mask = df.apply(lambda row: row.astype(str).str.contains(keyword, case=False).any(), axis=1)
                    filtered = df[mask]
                    if filtered.empty: 
                        return "관련 키워드 매칭 데이터 없음"
                    return filtered.head(max_rows).to_string()

                biz_summary = get_summary(df_biz, ind_str)
                keit_summary = get_summary(df_keit, tech_str if tech_str else ind_str)
                cert_summary = get_summary(df_cert, tech_str if tech_str else ind_str)
                
                # 3. AI 프롬프트에 융합 데이터 주입
                prompt = f"""
                현재 대상 기업은 '{loc_str}' 소재의 '{ind_str}' 업종이며, '{tech_str}' 기술을 다루고 있습니다.
                아래에 제공된 공공 빅데이터 요약본을 반드시 분석 근거로 활용하여 '생존율 정밀 진단 리포트'를 작성해 줘.

                [과거 공공데이터 기반 요약]
                - 전국사업체조사 관련 업종 통계: {biz_summary}
                - KEIT R&D 사업공고 트렌드: {keit_summary}
                - 기술개발제품 인증현황(벤치마킹 대상): {cert_summary}

                위 데이터를 근거로 하여 아래 4가지 목차에 맞춰 심사위원을 설득할 수 있는 수준의 전문적인 문장으로 작성할 것.
                1. 지역({loc_str}) 및 산업군 생존 환경 분석 (전국사업체조사 데이터 수치 인용)
                2. 주요 리스크 요인 3가지 (R&D 트렌드 기반의 기술, 자금, 시장 측면)
                3. 타사 기술인증 성공 사례를 벤치마킹한 리스크 극복 전략
                4. 향후 정부지원사업 확보 및 생존율 향상 방향성
                """
                
                try:
                    # 🚨 [수정] 챗봇 세션 대신 백그라운드 전용 모델로 일회성 호출
                    background_model = genai.GenerativeModel(model_name="gemini-2.5-flash")
                    response = background_model.generate_content(prompt)
                    st.session_state.survival_report = response.text
                except Exception as e:
                    st.error(f"AI 분석 중 오류가 발생했습니다: {e}")
                
    if st.session_state.survival_report:
        st.markdown(st.session_state.survival_report)

# ------- 지원 캘린더 설정 -----------
elif st.session_state.current_page == '지원 캘린더':
    from streamlit_calendar import calendar
    
    st.subheader("📅 나만의 맞춤형 지원 캘린더")
    st.caption("AI 매칭 메뉴에서 직접 '선택'하여 추가한 공고와 아래 폼으로 기입한 수동 일정만 깨끗하게 관리됩니다.")
    
    if 'calendar_events' not in st.session_state:
        st.session_state.calendar_events = []
        
    # --- 📝 1. 내가 원하는 일정 직접 기입 (수동 폼) 구역 ---
    with st.expander("📝 내가 원하는 수동 일정 직접 기입하기", expanded=False):
        with st.form("custom_schedule_form", clear_on_submit=True):
            custom_title = st.text_input("일정 및 사업명 입력", placeholder="예: 창업선도대학 대면평가 준비 또는 서류 제출 마감일")
            c_col1, c_col2 = st.columns(2)
            with c_col1:
                custom_start = st.date_input("시작 일자", datetime.now().date())
            with c_col2:
                custom_end = st.date_input("종료(마감) 일자", datetime.now().date())
            custom_url = st.text_input("참고용 상세 링크 URL (선택)", placeholder="https://...")
            
            submit_custom = st.form_submit_button("➕ 캘린더에 일정 강제 기입", use_container_width=True)
            if submit_custom:
                if not custom_title.strip():
                    st.warning("일정 이름을 정확히 입력해 주세요.")
                elif custom_start > custom_end:
                    st.error("시작 일자가 종료 일자보다 미래일 수 없습니다.")
                else:
                    # 마감일 당일까지 캘린더 칸이 꽉 차서 채워지도록 하루 더해줌
                    calc_end_str = (custom_end + timedelta(days=1)).strftime("%Y-%m-%d")
                    st.session_state.calendar_events.append({
                        "title": custom_title[:25] + "..." if len(custom_title) > 25 else custom_title,
                        "full_title": custom_title,
                        "start": custom_start.strftime("%Y-%m-%d"),
                        "end": calc_end_str,
                        "url": custom_url.strip()
                    })
                    st.success(f"'{custom_title}' 일정이 캘린더에 성공적으로 수동 등록되었습니다.")
                    st.rerun()

    st.markdown("---")
    
    # --- 📅 2. 선택/기입된 데이터 기반 클린 달력 렌더링 구역 ---
    if not st.session_state.calendar_events:
        st.info("💡 현재 캘린더에 등록된 마감 일정이 없습니다. 상단 폼을 통해 개인 일정을 기입하거나, 'AI 매칭' 탭에서 추천 공고를 선택해 등록해 보세요!")
    else:
        st.markdown("**📅 달력 기준월 이동 컨트롤러**")
        nav_col1, nav_col2, _ = st.columns([1, 1, 3])
        current_year = datetime.now().year
        current_month = datetime.now().month
        
        with nav_col1:
            target_year = st.selectbox("연도 선택", range(2020, 2031), index=range(2020, 2031).index(current_year), label_visibility="collapsed")
        with nav_col2:
            target_month = st.selectbox("월 선택", range(1, 13), index=current_month - 1, label_visibility="collapsed")
            
        target_date_str = f"{target_year}-{target_month:02d}-01"
        
        # 가공 데이터 매핑 (수동 등록 일정은 가시성을 위해 주황색, 공공 API 공고는 진청색 분리 적용)
        render_events = []
        for ev in st.session_state.calendar_events:
            is_official_api = bool(ev.get('url') and not ev.get('url').startswith('http') == False)
            render_events.append({
                "title": ev['title'],
                "start": ev['start'],
                "end": ev['end'],
                "backgroundColor": "#0e243a" if is_official_api else "#e67e22",
                "borderColor": "#0e243a" if is_official_api else "#e67e22"
            })
            
        calendar(events=render_events, options={
            "locale": "ko",
            "initialDate": target_date_str,
            "initialView": "dayGridMonth",
            "headerToolbar": {"left": "prev,next today", "center": "title", "right": "dayGridMonth,dayGridWeek"},
            "displayEventTime": False
        })
        
        # --- ⚙️ 3. 등록된 내부 마감 일정 리스트 조회 및 개별 삭제 관리 구역 ---
        st.markdown("---")
        st.subheader("🛠️ 현재 등록된 일정 관리 명단")
        
        management_list = []
        for idx, ev in enumerate(st.session_state.calendar_events):
            try:
                real_end_date = (datetime.strptime(ev['end'], "%Y-%m-%d").date() - timedelta(days=1)).strftime("%Y-%m-%d")
            except Exception:
                real_end_date = ev['end']
                
            management_list.append({
                "ID": idx,
                "구분": "🌐 공공 공고" if ev.get('url') else "📝 수동 일정",
                "일정/사업명": ev['full_title'],
                "마감일": real_end_date,
                "바로가기 링크": ev['url'] if str(ev['url']).startswith('http') else None
            })
            
        m_df = pd.DataFrame(management_list)
        st.dataframe(
            m_df[['구분', '일정/사업명', '마감일', '바로가기 링크']],
            use_container_width=True,
            column_config={
                "구분": st.column_config.TextColumn("구분", width="small"),
                "일정/사업명": st.column_config.TextColumn("일정/사업명", width="large"),
                "마감일": st.column_config.TextColumn("마감일", width="small"),
                "바로가기 링크": st.column_config.LinkColumn("바로가기 링크", display_text="🔗 바로가기", width="small")
            },
            hide_index=True
        )
        
        # 지저분해진 달력을 사용자가 능동적으로 제어할 수 있는 삭제 드롭다운 배포
        with st.expander("❌ 원치 않는 마감 일정 선택 삭제하기", expanded=False):
            select_options = [f"[{row['구분']}] {row['일정/사업명']}" for _, row in m_df.iterrows()]
            delete_choice = st.selectbox("삭제를 희망하는 일정을 고르세요:", options=select_options)
            
            if st.button("선택한 일정을 캘린더에서 완전 영구 삭제", type="secondary", use_container_width=True):
                matched_idx = m_df[m_df.apply(lambda r: f"[{r['구분']}] {r['일정/사업명']}" == delete_choice, axis=1)]["ID"].values[0]
                target_title = st.session_state.calendar_events[matched_idx]['full_title']
                
                # 세션 데이터에서 물리 제거
                st.session_state.calendar_events.pop(matched_idx)
                st.success(f"'{target_title}' 일정을 나침반 캘린더 데이터에서 완전히 안전하게 삭제했습니다.")
                st.rerun()

# ------- 보고서 생성 설정 -----------
elif st.session_state.current_page == '보고서 생성':
    st.header("📄 AI 자동 생성 보고서 보관함")
    st.caption("AI 매칭 메뉴에서 직접 선택하여 캘린더에 등록한 관심 공고들의 사업계획서 초안을 즉시 생성할 수 있습니다.")
    
    # 1. 🎯 [핵심] 캘린더에 저장된 관심 공고 리스트에서 사업명(full_title)만 쏙 뺴옵니다.
    saved_projects = []
    if 'calendar_events' in st.session_state and st.session_state.calendar_events:
        # 중복 방지 및 유효한 이름만 추출하여 리스트화
        saved_projects = [e['full_title'] for e in st.session_state.calendar_events if e.get('full_title')]
        
    # 2. 선택 옵션 구성 (선택된 공고가 있다면 최상단에 배치하고 알림을 띄웁니다)
    if saved_projects:
        options = saved_projects + ["직접 입력할게요 (선택)"]
        st.success(f"💡 현재 지원 캘린더와 연동되어 내가 찜한 공고 **{len(saved_projects)}건**을 바로 선택할 수 있습니다.")
    else:
        options = ["직접 입력할게요 (선택)"]
        st.info("💡 'AI 매칭' 탭에서 관심 공고를 선택해 캘린더에 담으면, 이곳에서 해당 공고를 번거로운 타이핑 없이 바로 선택하여 보고서를 작성할 수 있습니다.")
        
    # 3. 구성된 옵션으로 선택 박스 렌더링
    selected_business = st.selectbox(
        "작성할 지원사업명을 검색하거나 선택하세요:",
        options=options
    )
    
    if selected_business == "직접 입력할게요 (선택)":
        target_business = st.text_input("지원사업명을 직접 입력해주세요:")
    else:
        target_business = selected_business
        
    if st.button("사업계획서 초안 생성", type="primary"):
            if not target_business: 
                st.warning("사업명을 입력하거나 선택해주세요.")
            else:
                # 스피너 문구도 분석 내용에 맞게 살짝 변경해 줍니다.
                with st.spinner(f"'{target_business}' 지원 혜택 분석 및 실전용 사업계획서 작성 중..."):
                    
                    # 🎯 [핵심 수정] 실전용 사업계획서와 지원 이점을 동시에 도출하는 강력한 프롬프트
                    prompt = f"""
                    지원사업명: '{target_business}'
                    우리 기업 업종: {ind_str}
                    핵심 기술: {tech_str}
                    
                    당신은 수많은 스타트업을 합격시킨 정부지원사업 전문 최고 컨설턴트입니다.
                    단순한 범용 사업계획서가 아니라, 해당 사업에 '실제 지원'하기 위한 맞춤형 사업계획서 초안과, 이 사업에 선정되었을 때 우리 기업이 얻을 수 있는 '실질적인 이점'을 심층 분석해 줘.
                    
                    반드시 아래 목차에 맞추어 심사위원을 설득할 수 있는 전문적이고 구체적인 비즈니스 문구로 작성할 것.
                    
                    ---
                    
                    📌 [제 1부] 본 지원사업 선정 시 기대효과 (우리 기업의 실질적 이점)
                    1. 자금 조달 및 재무적 이점 (지원금 활용 가치)
                    2. 기술({tech_str}) 고도화 및 인프라 확보 측면
                    3. 시장 진출 및 레퍼런스(공공/민간) 확보 측면
                    
                    📝 [제 2부] {target_business} 실제 지원용 사업계획서 핵심 초안
                    1. 지원 사업 참여 동기 및 우리 기업 비전과의 적합성
                    2. 보유 기술의 차별성 및 시장 내 경쟁 우위
                    3. 타겟 시장 진입 및 비즈니스 모델(BM) 스케일업 전략
                    4. 지원금 활용 계획 및 구체적 향후 추진 일정
                    """
                    background_model = genai.GenerativeModel(model_name="gemini-2.5-flash")
                    response = background_model.generate_content(prompt)
                    st.session_state.report_content = response.text
                    st.session_state.current_target_business = target_business
                
    if "report_content" in st.session_state and "current_target_business" in st.session_state:
        st.markdown(st.session_state.report_content)
        
        clean_text = st.session_state.report_content
        clean_text = re.sub(r'\*+', '', clean_text)  
        clean_text = re.sub(r'#+\s*', '', clean_text) 
        clean_text = re.sub(r'`+', '', clean_text)    
        
        doc = Document()
        doc.add_heading(f"{st.session_state.current_target_business} 사업계획서", 0)
        doc.add_paragraph(clean_text)
        
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.divider()
        st.download_button(
            label="📥 기호 없이 깔끔한 워드(.docx) 다운로드", 
            data=bio.getvalue(), 
            file_name=f"{st.session_state.current_target_business}_사업계획서.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )